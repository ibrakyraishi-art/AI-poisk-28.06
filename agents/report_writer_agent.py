"""
Report Writer Agent — финальный отчёт в JSON + Word документ.

Порядок работы:
1. Агент получает в context_tasks структурированный вывод всех предыдущих агентов
   (ReviewData, NewsData, AnalystOutput) — CrewAI вставляет их JSON в промпт.
2. Агент формирует executive_summary и список рекомендаций.
3. Агент вызывает GenerateReportTool: передаёт полный ReportData JSON.
4. Инструмент создаёт .docx (python-docx) и загружает в Supabase Storage.
5. Агент возвращает ReportData (output_pydantic) — итоговый структурированный вывод.

Фолбэк: если Supabase Storage недоступен → .docx сохраняется в output/ локально.
Бакет "reports" должен быть создан в Supabase Storage до первого запуска.
"""
from __future__ import annotations

import io
import json
import os
from pathlib import Path

from crewai import Agent, Task
from crewai.tools import BaseTool
from docx import Document
from docx.shared import Pt

from config.llm_config import DEFAULTS, get_litellm_id
from models.schemas import AnalystOutput, ReportData, ReviewTheme
from tools.memory_search_tool import MemorySearchTool

_OUTPUT_DIR = Path(__file__).parent.parent / "output"


# ---------------------------------------------------------------------------
# docx builder
# ---------------------------------------------------------------------------

def _build_docx(report: ReportData) -> bytes:
    doc = Document()

    # Title
    doc.add_heading(f"{report.company} — Competitive Analysis Report", 0)
    doc.add_paragraph(f"Period: {report.period}  |  Generated: {report.generated_at}")

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(report.executive_summary)

    # User Review Themes
    doc.add_heading("User Review Themes", 1)
    for theme in report.review_themes:
        doc.add_heading(f"{theme.theme}  ({theme.sentiment})", 2)
        doc.add_paragraph(f"Mentioned in approximately {theme.review_count} reviews.")
        if theme.examples:
            doc.add_paragraph("Example quotes:")
            for quote in theme.examples:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f'"{quote}"').italic = True

    # SWOT Analysis
    doc.add_heading("SWOT Analysis", 1)
    swot = report.swot
    for label, items in [
        ("Strengths",     swot.strengths),
        ("Weaknesses",    swot.weaknesses),
        ("Opportunities", swot.opportunities),
        ("Threats",       swot.threats),
    ]:
        doc.add_heading(label, 2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    # Key Insights
    doc.add_heading("Key Insights", 1)
    for insight in swot.key_insights:
        doc.add_paragraph(insight, style="List Number")

    # Competitors (placeholder — filled in Stage 5)
    if report.competitors:
        doc.add_heading("Competitor Profiles", 1)
        for comp in report.competitors:
            doc.add_heading(comp.name, 2)
            if comp.rating:
                doc.add_paragraph(f"Rating: {comp.rating:.1f} / 5.0")
            if comp.price_model:
                doc.add_paragraph(f"Price model: {comp.price_model}")
            if comp.user_sentiment_summary:
                doc.add_paragraph(comp.user_sentiment_summary)

    # Recommendations
    doc.add_heading("Recommendations", 1)
    for i, rec in enumerate(report.recommendations, 1):
        doc.add_paragraph(f"{i}. {rec}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload_to_storage(docx_bytes: bytes, storage_path: str) -> str:
    """Upload docx to Supabase Storage bucket 'reports'. Returns the stored path."""
    from supabase import create_client  # noqa: PLC0415
    url = os.environ["SUPABASE_URL"]
    key = os.environ["SUPABASE_SERVICE_KEY"]
    client = create_client(url, key)
    client.storage.from_("reports").upload(
        storage_path,
        docx_bytes,
        file_options={
            "content-type": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
        },
    )
    return f"reports/{storage_path}"


# ---------------------------------------------------------------------------
# GenerateReportTool
# ---------------------------------------------------------------------------

class GenerateReportTool(BaseTool):
    """
    Takes a ReportData JSON string, builds a .docx file, and uploads it to
    Supabase Storage at reports/{run_id}/report.docx.
    Falls back to saving in the local output/ directory if Storage is unavailable.
    """
    name: str = "generate_report"
    description: str = (
        "Generate a Word document (.docx) report and save it to cloud storage. "
        "Input: a valid ReportData JSON string with ALL these fields: "
        "company (str), period (str), generated_at (ISO 8601, e.g. '2026-06-29T12:00:00Z'), "
        "executive_summary (str, 3-5 sentences), "
        "review_themes (list of objects: theme, sentiment, review_count, examples), "
        "swot (object: company, period, strengths, weaknesses, opportunities, threats, "
        "key_insights, confidence), "
        "competitors (empty list — will be filled in Stage 5), "
        "recommendations (list of str, 3-5 items). "
        "Returns: storage path of the saved report."
    )
    run_id: str
    user_id: str

    def _run(self, report_json: str) -> str:
        try:
            data = json.loads(report_json)
        except json.JSONDecodeError as exc:
            return f"Error: invalid JSON — {exc}. Fix the JSON and try again."

        try:
            report = ReportData.model_validate(data)
        except Exception as exc:
            return f"Error: ReportData validation failed — {exc}. Fix the data and try again."

        docx_bytes = _build_docx(report)
        storage_path = f"{self.run_id}/report.docx"

        try:
            saved_path = _upload_to_storage(docx_bytes, storage_path)
            return f"Report saved to Supabase Storage: {saved_path}"
        except Exception as exc:
            _OUTPUT_DIR.mkdir(exist_ok=True)
            local_path = _OUTPUT_DIR / f"report_{self.run_id}.docx"
            local_path.write_bytes(docx_bytes)
            return (
                f"Storage upload failed ({exc!r}). "
                f"Report saved locally: {local_path}"
            )


# ---------------------------------------------------------------------------
# Agent + Task factory functions
# ---------------------------------------------------------------------------

def create_report_writer_agent(
    *,
    run_id: str,
    user_id: str,
    model_key: str | None = None,
) -> Agent:
    model_key = model_key or DEFAULTS["report_writer_agent"]
    return Agent(
        role="Executive Report Writer",
        goal=(
            "Synthesise all collected data into a polished, management-ready competitive "
            "analysis report. Write a compelling executive summary, organise findings "
            "clearly, and produce concrete, prioritised recommendations."
        ),
        backstory=(
            "You are a senior management consultant who has written hundreds of competitive "
            "analysis reports for C-suite executives at tech companies. Your reports are "
            "renowned for combining sharp strategic insight with clear, concise prose. "
            "You never pad reports with vague statements — every sentence earns its place."
        ),
        tools=[
            GenerateReportTool(run_id=run_id, user_id=user_id),
            MemorySearchTool(run_id=run_id),
        ],
        llm=get_litellm_id(model_key),
        verbose=True,
        max_iter=6,
    )


def create_report_task(
    agent: Agent,
    *,
    company: str,
    period: str,
    context_tasks: list[Task],
) -> Task:
    return Task(
        description=(
            f"Write a comprehensive competitive analysis report for '{company}' "
            f"covering the past {period}.\n\n"
            "Steps:\n"
            "1. Review all previous agent outputs in your context:\n"
            "   ReviewData, NewsData, AnalystOutput, CompetitorFinderOutput.\n"
            "2. If you need additional quotes or facts, use memory_search.\n"
            "3. Draft the full report JSON:\n"
            "   - executive_summary: 3-5 sentences for senior management.\n"
            "     State: competitive position, top user concern, #1 opportunity, "
            "biggest competitive threat.\n"
            "   - review_themes: take directly from ReviewData context.\n"
            "   - swot: take directly from AnalystOutput context.\n"
            "   - competitors: take the full competitors list from "
            "CompetitorFinderOutput context. Include ALL profiles.\n"
            "   - recommendations: 3-5 items ordered by impact. Reference specific "
            "competitors where relevant.\n"
            "     Good: 'Fix Android GPS before RunKeeper markets the gap — "
            "it is the #1 driver of 1-star reviews and RunKeeper already highlights GPS accuracy.'\n"
            "     Bad: 'Improve the app.'\n"
            f"   - generated_at: today's date in ISO 8601 format.\n"
            "4. Call generate_report with the complete JSON. Fix validation errors "
            "and retry if needed.\n"
            "5. Output the final ReportData JSON."
        ),
        expected_output=(
            "A valid ReportData JSON with: company, period, generated_at (ISO 8601), "
            "executive_summary (str), review_themes (list of ReviewTheme), "
            "swot (AnalystOutput), competitors (list of CompetitorProfile from context), "
            "recommendations (list of str)."
        ),
        agent=agent,
        output_pydantic=ReportData,
        context=context_tasks,
    )
